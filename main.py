import docx  # 安装指令：pip install python-docx
import my_translate_func

def translate():
    #'''翻译'''

    # 获取文档对象
    doc = docx.Document('gen.docx')

    # 创建内存中的word文档对象
    new_doc = docx.Document()
    i=0
    # 遍历每一段文本
    for para in doc.paragraphs:
        # 写入新文件
        # new_doc.add_paragraph()
        nextex = para.text
        # 翻译
        print(nextex.lower())
        try:
            trans = my_translate_func.tran(nextex.lower())
        except:
            trans = ''
        # trans = youdao_translate(para.text)
        # trans = baidu_translate(para.text)
        # trans = my_translate_func.translate_you_dao2(nextex.lower())
        i += 1
        new_doc.add_paragraph(para.text+"\\\\"+trans)

    # 保存到本地文件
    new_doc.save('new-file-name.docx')

translate()
